"""
Unit & Integration Tests for Dynamic Prediction Report Generation (.docx).
Verifies that reports dynamically adapt to scenario, lead time, and fault injections.
"""

import os

import docx
from fastapi.testclient import TestClient

from backend.app import app
from replay.scenarios import ScenarioRunner
from reporting.docx_generator import create_3hour_prediction_docx


def test_create_3hour_prediction_docx_direct():
    runner = ScenarioRunner(config_path="config.yaml")
    snapshots = runner.run("config/scenarios/storm_01.yaml")
    out_file = os.path.join("outputs", "reports", "test_report_3hr.docx")

    res_path = create_3hour_prediction_docx(
        snapshots=snapshots,
        scenario_id="storm_01",
        output_path=out_file,
    )

    assert os.path.exists(res_path)
    assert os.path.getsize(res_path) > 1000  # Non-empty valid docx file


def test_download_docx_api_endpoint_baseline():
    client = TestClient(app)
    response = client.get("/api/reports/download-docx?scenario_id=storm_01")

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "attachment; filename=" in response.headers.get("content-disposition", "")
    assert len(response.content) > 1000


def test_download_docx_api_dynamic_lead_time_and_faults():
    client = TestClient(app)
    # Request dynamic report at t=45m with culvert blockage active
    response = client.get(
        "/api/reports/download-docx?scenario_id=storm_01&lead_time_minutes=45&fault_blockage=true"
    )

    assert response.status_code == 200
    assert "SafeSurge_Advisory_storm_01_t45m_clog.docx" in response.headers.get("content-disposition", "")
    
    # Save and inspect the document contents
    temp_path = os.path.join("outputs", "reports", "test_dynamic_t45.docx")
    with open(temp_path, "wb") as f:
        f.write(response.content)

    doc = docx.Document(temp_path)
    p_text = "\n".join([p.text for p in doc.paragraphs])
    table_text = "\n".join([cell.text for t in doc.tables for row in t.rows for cell in row.cells])
    full_text = p_text + "\n" + table_text
    
    # Verify dynamic content is present
    assert "+45 min" in full_text
    assert "Drain Clogged" in full_text or "Culvert E001" in full_text
    assert "Emergency Safe Navigation" in full_text
    assert "Street Grid Passability" in full_text
    assert "Field Sensor Telemetry" in full_text
    assert "Water Balance Invariant" in full_text


def test_download_docx_api_sensor_spike():
    client = TestClient(app)
    response = client.get(
        "/api/reports/download-docx?scenario_id=storm_01&lead_time_minutes=30&fault_spike=true"
    )

    assert response.status_code == 200
    assert "SafeSurge_Advisory_storm_01_t30m_surge.docx" in response.headers.get("content-disposition", "")

    temp_path = os.path.join("outputs", "reports", "test_dynamic_t30_surge.docx")
    with open(temp_path, "wb") as f:
        f.write(response.content)

    doc = docx.Document(temp_path)
    p_text = "\n".join([p.text for p in doc.paragraphs])
    table_text = "\n".join([cell.text for t in doc.tables for row in t.rows for cell in row.cells])
    full_text = p_text + "\n" + table_text

    assert "+30 min" in full_text
    assert "Water Surge" in full_text or "+90 cm" in full_text


def test_download_ml_report_endpoint():
    client = TestClient(app)
    response = client.get("/api/reports/download-ml-report")
    assert response.status_code == 200
    assert "AURA_FLOOD_ML_Training_and_Validation_Report.docx" in response.headers.get("content-disposition", "")
    assert len(response.content) > 10000

    # Parse and verify content
    temp_path = os.path.join("outputs", "reports", "test_ml_report_download.docx")
    with open(temp_path, "wb") as f:
        f.write(response.content)

    doc = docx.Document(temp_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Machine Learning" in full_text and "Validation" in full_text
    assert "XGBoost" in full_text or "AURA-FLOOD" in full_text
    assert len(doc.tables) >= 5



